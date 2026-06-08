import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import re

st.set_page_config(page_title="課表彙整系統", layout="wide")

# --- 核心替換函數 ---
def master_replace(doc_obj, old_text, new_text):
    if isinstance(new_text, (float, int)):
        new_val = str(int(new_text))
    else:
        new_val = str(new_text) if (new_text and str(new_text).strip() != "") else ""
    targets = list(doc_obj.paragraphs)
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                targets.extend(cell.paragraphs)
    for p in targets:
        if old_text in p.text:
            full_text = "".join([run.text for run in p.runs])
            updated_text = full_text.replace(old_text, new_val)
            for i, run in enumerate(p.runs):
                run.text = updated_text if i == 0 else ""

# --- 讀取內建樣板函數 ---
def load_default_template(file_name):
    try:
        with open(file_name, "rb") as f:
            return f.read()
    except FileNotFoundError:
        return None

# --- 側邊欄：徹底簡化 ---
with st.sidebar:
    st.header("⚙️ 資料管理")
    if st.button("🧹 清空重置系統"):
        for key in list(st.session_state.keys()): del st.session_state[key]
        st.rerun()

    # --- 📥 僅保留資料範本下載 ---
    st.divider()
    st.subheader("📥 範本下載")
    data_templates = {
        "1. 配課表範本": "配課表.xlsx",
        "2. 課表範本": "課表.xlsx",
        "3. 教師排序表範本": "教師排序表.xlsx"
    }
    for label, file_name in data_templates.items():
        try:
            with open(file_name, "rb") as f:
                st.download_button(label=f"{label}", data=f, file_name=file_name, key=f"dl_{file_name}")
        except FileNotFoundError:
            st.caption(f"⚠️ 找不到 {file_name}")
    st.divider()

    # --- 📤 僅保留資料上傳 ---
    st.subheader("📤 上傳資料檔")
    f_assign = st.file_uploader("1. 上傳【配課表】", type=["xlsx", "csv"])
    f_time = st.file_uploader("2. 上傳【課表】", type=["xlsx", "csv"])
    f_sort = st.file_uploader("3. 上傳【教師排序暨時數表】", type=["xlsx", "csv"])
    
    if f_assign and f_time and st.button("🚀 執行整合"):
        # 自動抓取後台 Word 樣板
        class_temp = load_default_template("班級樣板.docx")
        teacher_temp = load_default_template("教師樣板.docx")
        
        if not class_temp or not teacher_temp:
            st.error("❌ 系統錯誤：後台找不到「班級樣板.docx」或「教師樣板.docx」，請確認 GitHub 檔案。")
        else:
            with st.spinner("同步內建樣板與解析資料中..."):
                df_assign = pd.read_csv(f_assign) if f_assign.name.endswith('.csv') else pd.read_excel(f_assign)
                df_time = pd.read_csv(f_time) if f_time.name.endswith('.csv') else pd.read_excel(f_time)
                
                # 存入 Session 供後續下載使用
                st.session_state.class_template = class_temp
                st.session_state.teacher_template = teacher_temp

                # 1. 解析配課 (多師共課)
                assign_lookup, all_teachers_db, tutors = [], set(), {}
                for _, row in df_assign.iterrows():
                    c, s, t_raw = str(row['班級']).strip(), str(row['科目']).strip(), str(row['教師']).strip()
                    t_list = [name.strip() for name in t_raw.split('/')]
                    for t in t_list:
                        if t and t != "nan":
                            assign_lookup.append({'c': c, 's': s, 't': t})
                            all_teachers_db.add(t)
                    if s == "班級": tutors[c] = t_raw

                # 2. 教師排序與時數
                ordered_teachers, base_hours, all_teachers_list = [], {}, list(all_teachers_db)
                if f_sort:
                    df_s = pd.read_csv(f_sort) if f_sort.name.endswith('.csv') else pd.read_excel(f_sort)
                    for _, s_row in df_s.iterrows():
                        t_name = str(s_row.iloc[0]).strip()
                        if t_name in all_teachers_list:
                            ordered_teachers.append(t_name)
                            try: base_hours[t_name] = int(s_row.iloc[1])
                            except: base_hours[t_name] = 0
                    for t in all_teachers_list:
                        if t not in ordered_teachers: ordered_teachers.append(t); base_hours[t] = 0
                else:
                    ordered_teachers = sorted(all_teachers_list)
                    base_hours = {t: 0 for t in ordered_teachers}

                # 3. 解析課表
                class_data, teacher_data, total_counts = {}, {}, {}
                day_map = {"一":1,"二":2,"三":3,"四":4,"五":5,"週一":1,"週二":2,"週三":3,"週四":4,"週五":5}
                for _, row in df_time.iterrows():
                    c_raw, s_raw = str(row['班級']).strip(), str(row['科目']).strip()
                    d, p_match = day_map.get(str(row['星期']).strip(), 0), re.search(r'\d+', str(row['節次']))
                    if not (p_match and d > 0): continue
                    p = int(p_match.group())

                    # 修正：先檢查科目是否為空，如果是空的，就不去配課表找老師
                    if not s_raw or s_raw == "nan" or s_raw == "":
                        display_t = ""
                        s_raw = ""  # 讓科目也保持空白，不要出現 nan
                    else:
                        curr_t_list = [item['t'] for item in assign_lookup if item['c'] == c_raw and item['s'] == s_raw]
                        display_t = "/".join(curr_t_list) if curr_t_list else "未知教師"
                    
                    if c_raw not in class_data: class_data[c_raw] = {}
                    class_data[c_raw][(d, p)] = {"subj": s_raw, "teacher": display_t}
                    
                    for t in curr_t_list:
                        if t not in teacher_data: teacher_data[t] = {}
                        teacher_data[t][(d, p)] = {"subj": s_raw, "class": c_raw}
                        total_counts[t] = total_counts.get(t, 0) + 1

                st.session_state.update({
                    "class_data": class_data, "teacher_data": teacher_data, "tutors_map": tutors,
                    "base_hours": base_hours, "total_counts": total_counts, "ordered_teachers": ordered_teachers,
                    "sel_class": sorted(list(class_data.keys()))[0], "sel_teacher": ordered_teachers[0]
                })
                st.rerun()

# --- 主介面與預覽 (代碼同前，保留班級預覽顯示科目(老師)功能) ---
if 'class_data' in st.session_state:
    tab1, tab2 = st.tabs(["🏫 班級課表", "👩‍🏫 教師課表"])

    with tab1:
        classes = sorted(list(st.session_state.class_data.keys()))
        curr_c = st.session_state.get('sel_class', classes[0])
        col1, col2, col3 = st.columns([1, 2, 1])
        if col1.button("⬅️ 上一班"):
            st.session_state.sel_class = classes[(classes.index(curr_c) - 1) % len(classes)]; st.rerun()
        if col3.button("下一班 ➡️"):
            st.session_state.sel_class = classes[(classes.index(curr_c) + 1) % len(classes)]; st.rerun()
        with col2: st.session_state.sel_class = st.selectbox("選取班級", classes, index=classes.index(curr_c))
        
        target_c = st.session_state.sel_class
        st.info(f"📍 班級：{target_c} | 導師：{st.session_state.tutors_map.get(target_c, '未設定')}")
        
        c_preview = []
        for p in range(1, 9):
            row = {"節次": f"第 {p} 節"}
            for d in range(1, 6):
                info = st.session_state.class_data[target_c].get((d,p))
                row[f"週{d}"] = f"{info['subj']}\n({info['teacher']})" if info else ""
            c_preview.append(row)
        st.table(pd.DataFrame(c_preview))

        bc1, bc2 = st.columns(2)
        with bc1:
            if st.button(f"📥 下載 {target_c} 課表"):
                doc = Document(BytesIO(st.session_state.class_template))
                # [關鍵修正 1]：執行班級與導師的替換
                master_replace(doc, "{{CLASS}}", target_c)
                tutor_name = st.session_state.tutors_map.get(target_c, "未設定")
                master_replace(doc, "{{TUTOR}}", tutor_name) 
                
                for d, p in [(d,p) for d in range(1,6) for p in range(1,9)]:
                    v = st.session_state.class_data[target_c].get((d,p), {"subj":"","teacher":""})
                    master_replace(doc, f"{{{{SD{d}P{p}}}}}", v['subj'])
                    master_replace(doc, f"{{{{TD{d}P{p}}}}}", v['teacher'])
                
                buf = BytesIO()
                doc.save(buf)
                st.download_button(f"💾 儲存 {target_c} 課表", buf.getvalue(), f"{target_c}_班級課表.docx")
        with bc2:
            sel_c_batch = st.multiselect("勾選批次合併", classes, default=classes)
            if st.button("🚀 執行班級合併列印"):
                main_doc = None
                for i, cname in enumerate(sel_c_batch):
                    tmp = Document(BytesIO(st.session_state.class_template)); master_replace(tmp, "{{CLASS}}", cname)
                    
                    master_replace(tmp, "{{CLASS}}", cname)
                    master_replace(tmp, "{{TUTOR}}", st.session_state.tutors_map.get(cname, "未設定"))
                    
                    for d, p in [(d,p) for d in range(1,6) for p in range(1,9)]:
                        v = st.session_state.class_data[cname].get((d,p), {"subj":"","teacher":""})
                        master_replace(tmp, f"{{{{SD{d}P{p}}}}}", v['subj']); master_replace(tmp, f"{{{{TD{d}P{p}}}}}", v['teacher'])
                    if i == 0: main_doc = tmp
                    else: 
                        for el in tmp.element.body: main_doc.element.body.append(el)
                if main_doc:
                    buf = BytesIO(); main_doc.save(buf); st.download_button("💾 下載班級彙整檔", buf.getvalue(), "全校班級課表.docx")

    with tab2:
        # (教師標籤頁同樣保留原本強大的預覽與下載功能)
        teachers = st.session_state.ordered_teachers
        curr_t = st.session_state.get('sel_teacher', teachers[0])
        colt1, colt2, colt3 = st.columns([1, 2, 1])
        if colt1.button("⬅️ 前一位"):
            st.session_state.sel_teacher = teachers[(teachers.index(curr_t) - 1) % len(teachers)]; st.rerun()
        if colt3.button("下一位 ➡️"):
            st.session_state.sel_teacher = teachers[(teachers.index(curr_t) + 1) % len(teachers)]; st.rerun()
        with colt2: st.session_state.sel_teacher = st.selectbox("跳轉教師", teachers, index=teachers.index(curr_t))

        target_t = st.session_state.sel_teacher
        base, total = int(st.session_state.base_hours.get(target_t, 0)), int(st.session_state.total_counts.get(target_t, 0))
        m1, m2, m3 = st.columns(3); m1.metric("應授時數", f"{base} 節"); m2.metric("教學總時數", f"{total} 節"); m3.metric("兼代課時數", f"{total-base} 節")
        
        t_prev = [{"節次": f"第 {p} 節", **{f"週{d}": f"{st.session_state.teacher_data[target_t].get((d,p),{}).get('class','')} {st.session_state.teacher_data[target_t].get((d,p),{}).get('subj','')}".strip() for d in range(1,6)}} for p in range(1,9)]
        st.table(pd.DataFrame(t_prev))

        bt1, bt2 = st.columns(2)
        with bt1:
            if st.button(f"📥 下載 {target_t} 課表"):
                doc = Document(BytesIO(st.session_state.teacher_template))
                master_replace(doc, "{{TEACHER}}", target_t); master_replace(doc, "{{BASE}}", base)
                master_replace(doc, "{{TOTAL}}", total); master_replace(doc, "{{EXTRA}}", total-base)
                for d, p in [(d,p) for d in range(1,6) for p in range(1,9)]:
                    v = st.session_state.teacher_data[target_t].get((d,p), {"subj":"","class":""})
                    master_replace(doc, f"{{{{CD{d}P{p}}}}}", v['class']); master_replace(doc, f"{{{{SD{d}P{p}}}}}", v['subj'])
                buf = BytesIO(); doc.save(buf); st.download_button(f"💾 儲存 {target_t} 課表", buf.getvalue(), f"{target_t}_教師課表.docx")
        with bt2:
            sel_t_batch = st.multiselect("批次合併教師", teachers, default=teachers)
            if st.button("🚀 執行教師合併列印"):
                main_doc = None
                for i, tname in enumerate(sel_t_batch):
                    tb, tt = int(st.session_state.base_hours.get(tname, 0)), int(st.session_state.total_counts.get(tname, 0))
                    tmp = Document(BytesIO(st.session_state.teacher_template))
                    master_replace(tmp, "{{TEACHER}}", tname); master_replace(tmp, "{{BASE}}", tb)
                    master_replace(tmp, "{{TOTAL}}", tt); master_replace(tmp, "{{EXTRA}}", tt-tb)
                    for d, p in [(d,p) for d in range(1,6) for p in range(1,9)]:
                        v = st.session_state.teacher_data[tname].get((d,p), {"subj":"","class":""})
                        master_replace(tmp, f"{{{{CD{d}P{p}}}}}", v['class']); master_replace(tmp, f"{{{{SD{d}P{p}}}}}", v['subj'])
                    if i == 0: main_doc = tmp
                    else: 
                        for el in tmp.element.body: main_doc.element.body.append(el)
                if main_doc:
                    buf = BytesIO(); main_doc.save(buf); st.download_button("💾 下載教師彙整檔", buf.getvalue(), "全校教師課表_彙整.docx")
else:
    st.info("👋 請上傳資料檔並點擊執行整合。")










